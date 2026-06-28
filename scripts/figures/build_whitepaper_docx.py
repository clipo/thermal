#!/usr/bin/env python3
"""Compile docs/WHITEPAPER.md into docs/WHITEPAPER.docx with figures embedded
inline at the points where they appear in the markdown.

Handles: ATX headings, inline images (![alt](path)), bold figure-caption
paragraphs (**Figure N.** ...), markdown pipe tables, and ordinary
paragraphs (hard-wrapped source lines are collapsed into one Word
paragraph). Inline emphasis and code markers are stripped to plain text.

Run: python scripts/figures/build_whitepaper_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

THERMAL = Path(__file__).resolve().parent.parent.parent
DOCS = THERMAL / "docs"
PAPER_MD = DOCS / "WHITEPAPER.md"
OUT_DOCX = DOCS / "WHITEPAPER.docx"

IMG_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")


def strip_inline(text: str) -> str:
    """Remove markdown emphasis/code markers, keep the text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def add_para(doc, text):
    p = doc.add_paragraph()
    p.add_run(strip_inline(text)).font.size = Pt(11)
    return p


def add_caption(doc, text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(strip_inline(text))
    run.font.size = Pt(9)
    run.font.italic = True


def add_image(doc, rel_path: str):
    img = (DOCS / rel_path).resolve()
    if not img.exists():
        print(f"  MISSING image: {img}")
        doc.add_paragraph(f"[image missing: {rel_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Inches(6.0))
    print(f"  embedded {img.name}")


def add_table(doc, rows: list[list[str]]):
    """rows includes header; the markdown separator row is already removed."""
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Light Grid Accent 1"
    for ri, cells in enumerate(rows):
        wrow = table.add_row().cells
        for ci in range(ncol):
            text = strip_inline(cells[ci]) if ci < len(cells) else ""
            wrow[ci].text = text
            for para in wrow[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if ri == 0:
                        run.font.bold = True


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_sep_row(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", line)) and "-" in line


def main():
    if not PAPER_MD.exists():
        raise SystemExit(f"missing {PAPER_MD}")
    lines = PAPER_MD.read_text().splitlines()

    doc = Document()
    sec = doc.sections[0]
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(1.0))

    buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_para():
        nonlocal buf
        if buf:
            text = re.sub(r"\s+", " ", " ".join(b.strip() for b in buf)).strip()
            if text:
                add_para(doc, text)
            buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

    for line in lines:
        stripped = line.strip()

        # Table accumulation
        if stripped.startswith("|"):
            flush_para()
            if is_sep_row(stripped):
                continue
            table_buf.append(split_table_row(stripped))
            continue
        else:
            flush_table()

        m = IMG_RE.match(stripped)
        if m:
            flush_para()
            add_image(doc, m.group("path"))
            continue

        if stripped.startswith("#"):
            flush_para()
            level = len(stripped) - len(stripped.lstrip("#"))
            heading = stripped[level:].strip()
            doc.add_heading(strip_inline(heading), level=min(level, 4))
            continue

        if stripped.startswith("---"):
            flush_para()
            continue

        if stripped == "":
            flush_para()
            continue

        # Figure caption paragraph: render italic/centered as its own block.
        if stripped.startswith("**Figure"):
            flush_para()
            add_caption(doc, stripped)
            continue

        # Ordered/unordered list item: emit each item as its own paragraph.
        # Continuation (wrapped) lines fall through to buf and collapse in.
        if re.match(r"^(\d+\.|[-*])\s", stripped):
            flush_para()
            buf.append(stripped)
            continue

        buf.append(line)

    flush_para()
    flush_table()
    doc.save(str(OUT_DOCX))
    print(f"\nWrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
