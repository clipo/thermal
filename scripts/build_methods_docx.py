#!/usr/bin/env python3
"""Convert PAPER_METHODS.md into a Word .docx with embedded site-closeup
figures, suitable for the Methods section of an academic paper.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


THERMAL = Path(__file__).resolve().parent.parent
PAPER_MD = THERMAL / "PAPER_METHODS.md"
OUT_DOCX = THERMAL / "PAPER_METHODS.docx"
FIG_DIR = THERMAL / "sgd_output" / "figures"
CLOSEUPS = FIG_DIR / "closeups"

# Figure mapping: caption number → (image path, caption text)
FIGURES = [
    ("Figure 1",
     FIG_DIR / "rapa_nui_overview_coastal.png",
     "Island-wide SGD distribution. Coast-anchored SGD plumes detected "
     "across 29 thermal-drone flights of Rapa Nui (June 2023 and January "
     "2024). Each point represents one plume polygon; size and colour "
     "scale with the polygon's Σ_anomaly (m² × °C of integrated cold-"
     "anomaly content). Background shaded blocks show flight footprints. "
     "The full-resolution SVG and KML versions allow toggling individual "
     "flights and inspecting polygon attributes."),
    ("Figure 2",
     CLOSEUPS / "vaihu_full_Vaihu_Harbor_closeup.png",
     "Vaihu Harbor (textbook SGD reference site). Coast-anchored plumes "
     "overlaid on Esri WorldImagery satellite tile. Eighteen discrete "
     "plumes traced along the bay shoreline, with sources at the rocky "
     "cape and inner harbour where freshwater is known to emerge through "
     "collapsed lava-tube outlets. The 0.3 °C anomaly contour (red "
     "dashed) shows the broader plume envelope."),
    ("Figure 3",
     CLOSEUPS / "june2023_23_june_23_tongariki_flights_Hanga_Nui_(Ahu_Tongariki)_closeup.png",
     "Hanga Nui at Ahu Tongariki. Five discrete plumes detected in the "
     "bay where the famous moai platform is sited. Bay geometry typical "
     "of Rapa Nui's SGD-active sites: low-elevation beach/rocks, "
     "sheltered from the open Pacific, with Poike's volcanic slopes "
     "inland."),
    ("Figure 4",
     CLOSEUPS / "flight8_hekii_west_Hekii_West_closeup.png",
     "Hekii West. Twelve discrete plumes along a known SGD hot zone on "
     "the east coast. Strong cold-anomaly signal (peak ~1 °C below "
     "ambient) typical of strong-flow lava-tube outlets."),
    ("Figure 5",
     CLOSEUPS / "flight10_anakena_to_west_Anakena_Bay_closeup.png",
     "Anakena Bay. Four plumes detected in the small sandy-beach bay. "
     "The relatively low Σ_anomaly here is consistent with Anakena's "
     "geology — small drainage basin upstream, modest groundwater "
     "throughput."),
    ("Figure 6",
     CLOSEUPS / "flight11_hivahiva_to_hangapiko_Hivahiva-Hangapiko_closeup.png",
     "Hivahiva-Hangapiko (south coast). Ten plumes along the south "
     "coast bay system with strong cold signal at the rocky inshore "
     "zone."),
    ("Figure 7",
     CLOSEUPS / "june2023_2_july_23_poike_3_Poike_(cliff_coast_control)_closeup.png",
     "Poike (cliff-coast control). The cliff-zone filter based on SRTM "
     "elevation correctly suppresses the spurious cliff-shadow polygons "
     "that would otherwise dominate this flight. Surviving plumes are "
     "at the surf zone immediately offshore of the cliff base, where "
     "some real cold-water signal may exist (or may still represent "
     "residual misprojection — flagged as a known limitation)."),
]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_figure(doc, label, image_path, caption):
    if not image_path.exists():
        print(f"  ✗ missing image: {image_path}")
        doc.add_paragraph(f"[{label}: image missing — {image_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"{label}. {caption}")
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    print(f"  ✓ embedded {label} ({image_path.name})")


def _flush_para(doc, buf):
    """Emit buffered lines as ONE Word paragraph — collapse hard-wrapped
    markdown line breaks into spaces so paragraphs flow naturally."""
    if not buf:
        return
    text = " ".join(line.strip() for line in buf).strip()
    text = re.sub(r"\s+", " ", text)
    add_para(doc, text)


def parse_markdown_into_doc(md_text: str, doc: Document):
    """Walk the markdown, emit headings/paragraphs into doc.
    Skip the trailing 'Figure captions' section (figures are embedded
    inline in the FIGURES list later by main()). Hard-wrapped lines
    in source markdown are collapsed into single Word paragraphs."""
    body = md_text.split("## Figure captions")[0]
    buf: list[str] = []
    for line in body.splitlines():
        if line.startswith("# "):
            _flush_para(doc, buf); buf = []
            add_heading(doc, line[2:].strip(), level=0)
        elif line.startswith("## "):
            _flush_para(doc, buf); buf = []
            add_heading(doc, line[3:].strip(), level=1)
        elif line.startswith("### "):
            _flush_para(doc, buf); buf = []
            add_heading(doc, line[4:].strip(), level=2)
        elif line.startswith("---"):
            continue
        elif line.strip() == "":
            _flush_para(doc, buf); buf = []
        else:
            buf.append(line)
    _flush_para(doc, buf)


def main():
    if not PAPER_MD.exists():
        raise SystemExit(f"missing {PAPER_MD}")

    md_text = PAPER_MD.read_text()
    doc = Document()
    # Page setup: standard letter, 1" margins
    sec = doc.sections[0]
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    parse_markdown_into_doc(md_text, doc)

    # Insert figures section as a block at the end
    add_heading(doc, "Figures", level=1)
    for label, path, caption in FIGURES:
        add_figure(doc, label, path, caption)
        doc.add_paragraph()  # spacing

    doc.save(str(OUT_DOCX))
    print(f"\nWrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
